from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


def normalize_text(text):
    # If accidentally list or tuple is passed, fix it
    if isinstance(text, list):
        return ", ".join(text)
    return str(text)


def create_docx(resume):
    doc = Document()

    doc.add_heading(resume["name"], level=0)
    doc.add_paragraph(resume["role"])

    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(normalize_text(resume["summary"]))

    doc.add_heading("TECHNICAL SKILLS", level=1)
    doc.add_paragraph(normalize_text(resume["skills"]))

    doc.add_heading("PROJECT EXPERIENCE", level=1)
    for p in resume["projects"]:
        doc.add_paragraph(f"- {normalize_text(p)}")

    doc.add_heading("WORK EXPERIENCE", level=1)
    doc.add_paragraph(normalize_text(resume["experience"]))

    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph(normalize_text(resume["education"]))

    doc.save("resume.docx")


def create_pdf(resume):
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate("resume.pdf")

    story = []

    def add(title, content):
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        story.append(Paragraph(normalize_text(content), styles["Normal"]))

    story.append(Paragraph(f"<b>{resume['name']}</b>", styles["Title"]))
    story.append(Paragraph(resume["role"], styles["Normal"]))

    add("PROFESSIONAL SUMMARY", resume["summary"])
    add("TECHNICAL SKILLS", resume["skills"])

    story.append(Paragraph("<b>PROJECT EXPERIENCE</b>", styles["Heading2"]))
    for p in resume["projects"]:
        story.append(Paragraph("- " + normalize_text(p), styles["Normal"]))

    add("WORK EXPERIENCE", resume["experience"])
    add("EDUCATION", resume["education"])

    pdf.build(story)
